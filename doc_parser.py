"""
Document parser — extracts text and tables from documents.

Uses Python libraries (PyMuPDF, python-docx, openpyxl) for text-based formats
and Azure AI Document Intelligence OCR for scanned documents and images.

Performance: files are parsed concurrently during upload.
"""

import io
import importlib
import logging
import os

logger = logging.getLogger("vigil.doc_parser")

IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "tiff", "tif", "bmp"}


def _env_int(name: str, default: int) -> int:
    try:
        return int(os.getenv(name, str(default)))
    except (TypeError, ValueError):
        return default


def _env_float(name: str, default: float) -> float:
    try:
        return float(os.getenv(name, str(default)))
    except (TypeError, ValueError):
        return default


PDF_OCR_MODE = os.getenv("PDF_OCR_MODE", "auto").strip().lower()
PDF_MIN_EMBEDDED_TEXT_CHARS = max(20, _env_int("PDF_MIN_EMBEDDED_TEXT_CHARS", 120))
PDF_LOW_TEXT_PAGE_CHARS = max(5, _env_int("PDF_LOW_TEXT_PAGE_CHARS", 35))
PDF_LOW_TEXT_PAGE_RATIO = min(1.0, max(0.0, _env_float("PDF_LOW_TEXT_PAGE_RATIO", 0.4)))
PDF_OCR_MIN_GAIN_FACTOR = max(1.0, _env_float("PDF_OCR_MIN_GAIN_FACTOR", 1.15))


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_document(filename: str, content_bytes: bytes) -> str:
    """Parse a document file and return its text content.

    Uses Python libraries for text-based formats and Azure AI Document
    Intelligence OCR for scanned documents and images.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "txt":
        return content_bytes.decode("utf-8", errors="replace")
    elif ext == "pdf":
        return _parse_pdf(content_bytes)
    elif ext == "docx":
        return _parse_docx(content_bytes)
    elif ext in ("xlsx", "xls"):
        return _parse_xlsx(content_bytes)
    elif ext in IMAGE_EXTENSIONS:
        return _ocr_with_document_intelligence(content_bytes, filename)
    else:
        # Try as plain text
        return content_bytes.decode("utf-8", errors="replace")


# --- PDF (PyMuPDF / pdfplumber) ---

def _parse_pdf(data: bytes) -> str:
    """Extract text from PDF and use OCR heuristically for low-text/scanned files.

    OCR behavior is controlled by PDF_OCR_MODE:
    - off: never use OCR for PDFs
    - auto: OCR only when embedded text looks sparse
    - hybrid: OCR on sparse PDFs and merge OCR with embedded text
    - force: always run OCR (fallback to embedded text if OCR fails)
    """
    text = ""
    page_count = 0
    low_text_pages = 0

    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype="pdf")
        parts = []
        page_count = len(doc)
        for page in doc:
            page_text = page.get_text() or ""
            parts.append(page_text)
            if len(page_text.strip()) < PDF_LOW_TEXT_PAGE_CHARS:
                low_text_pages += 1
        doc.close()
        text = "\n".join(parts)
    except ImportError:
        logger.warning("PyMuPDF not installed — trying pdfplumber")
    except Exception as exc:
        logger.error("PyMuPDF failed: %s", exc)

    if not text.strip():
        try:
            pdfplumber = importlib.import_module("pdfplumber")
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                parts = []
                page_count = len(pdf.pages)
                low_text_pages = 0
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    parts.append(page_text)
                    if len(page_text.strip()) < PDF_LOW_TEXT_PAGE_CHARS:
                        low_text_pages += 1
                text = "\n".join(parts)
        except ImportError:
            pass
        except Exception as exc:
            logger.error("pdfplumber failed: %s", exc)

    mode = PDF_OCR_MODE if PDF_OCR_MODE in {"off", "auto", "hybrid", "force"} else "auto"
    stripped = text.strip()
    text_len = len(stripped)
    low_text_ratio = (low_text_pages / page_count) if page_count else 0.0

    should_try_ocr = False
    if mode == "force":
        should_try_ocr = True
    elif mode in {"auto", "hybrid"}:
        should_try_ocr = (
            text_len < PDF_MIN_EMBEDDED_TEXT_CHARS or
            (page_count > 0 and low_text_ratio >= PDF_LOW_TEXT_PAGE_RATIO)
        )

    if should_try_ocr:
        logger.info(
            "PDF OCR trigger (mode=%s, text=%d chars, low-text pages=%d/%d)",
            mode,
            text_len,
            low_text_pages,
            page_count,
        )
        ocr_text = _ocr_with_document_intelligence(data, "document.pdf")
        if _is_valid_ocr_text(ocr_text):
            ocr_len = len(ocr_text.strip())
            if mode == "force":
                return ocr_text
            if mode == "hybrid":
                return _merge_pdf_text_sources(text, ocr_text)
            if text_len == 0 or ocr_len > int(text_len * PDF_OCR_MIN_GAIN_FACTOR):
                logger.info("Using OCR text for PDF (embedded=%d, ocr=%d chars)", text_len, ocr_len)
                return ocr_text
            logger.info("Keeping embedded PDF text (embedded=%d, ocr=%d chars)", text_len, ocr_len)
        elif mode == "force" and stripped:
            logger.warning("Forced OCR did not return usable text, falling back to embedded PDF text")

    return text if text.strip() else "[No text could be extracted from this PDF]"


def _is_valid_ocr_text(ocr_text: str) -> bool:
    return bool(ocr_text and ocr_text.strip() and not ocr_text.startswith("["))


def _merge_pdf_text_sources(embedded_text: str, ocr_text: str) -> str:
    """Merge two text sources while avoiding obvious duplication."""
    embedded = (embedded_text or "").strip()
    ocr = (ocr_text or "").strip()

    if not embedded:
        return ocr
    if not ocr:
        return embedded
    if embedded in ocr:
        return ocr
    if ocr in embedded:
        return embedded
    return f"{embedded}\n\n[OCR supplement]\n{ocr}"


# --- DOCX (python-docx) ---

def _parse_docx(data: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except ImportError:
        return "[DOCX parsing requires python-docx]"
    except Exception as exc:
        logger.error("DOCX parsing failed: %s", exc)
        return f"[DOCX parsing error: {exc}]"


# --- XLSX (openpyxl) ---

def _parse_xlsx(data: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                vals = [str(c) if c is not None else "" for c in row]
                if any(v.strip() for v in vals):
                    lines.append(" | ".join(vals))
        return "\n".join(lines)
    except ImportError:
        return "[XLSX parsing requires openpyxl]"
    except Exception as exc:
        logger.error("XLSX parsing failed: %s", exc)
        return f"[XLSX parsing error: {exc}]"


# ---------------------------------------------------------------------------
# Azure AI Document Intelligence OCR (scanned PDFs and images)
# ---------------------------------------------------------------------------

# Cached OCR client (initialized lazily, thread-safe for concurrent uploads)
_ocr_client = None


def _get_ocr_client():
    """Return a cached DocumentIntelligenceClient, or None if not configured."""
    global _ocr_client
    if _ocr_client is not None:
        return _ocr_client

    try:
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.identity import DefaultAzureCredential
        from foundry_client import get_cognitive_endpoint

        endpoint = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT", "").rstrip("/")
        if not endpoint:
            try:
                endpoint = get_cognitive_endpoint()
            except ValueError:
                return None
        if not endpoint:
            return None

        _ocr_client = DocumentIntelligenceClient(
            endpoint=endpoint, credential=DefaultAzureCredential()
        )
        logger.info("OCR client initialized for %s", endpoint)
        return _ocr_client
    except ImportError:
        logger.warning("azure-ai-documentintelligence not installed")
        return None


def _ocr_with_document_intelligence(data: bytes, filename: str) -> str:
    """Extract text from a scanned document or image using Azure AI Document Intelligence."""
    try:
        from azure.ai.documentintelligence.models import AnalyzeDocumentRequest

        client = _get_ocr_client()
        if not client:
            return "[OCR requires AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT or FOUNDRY_PROJECT_ENDPOINT to be set]"

        poller = client.begin_analyze_document(
            "prebuilt-layout",
            AnalyzeDocumentRequest(bytes_source=data),
        )
        result = poller.result()

        # prebuilt-layout returns structured content with tables as markdown
        if hasattr(result, "content") and result.content:
            text = result.content
            logger.info("Azure OCR (layout) extracted %d chars from %s", len(text), filename)
            return text if text.strip() else "[OCR completed but no text was detected]"

        # Fallback: extract line-by-line from pages
        pages_text = []
        for page in result.pages:
            lines = [line.content for line in (page.lines or [])]
            pages_text.append("\n".join(lines))

        text = "\n\n".join(pages_text)
        logger.info("Azure OCR (layout/lines) extracted %d chars from %s", len(text), filename)
        return text if text.strip() else "[OCR completed but no text was detected]"

    except ImportError:
        logger.warning("azure-ai-documentintelligence not installed")
        return "[OCR requires azure-ai-documentintelligence package]"
    except Exception as e:
        logger.error("Azure Document Intelligence OCR failed: %s", e)
        return f"[OCR failed: {e}]"
